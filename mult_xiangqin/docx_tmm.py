import requests
from lxml import etree

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def spider_page():
    document = Document()

    # word页面计数
    count = 1
    for page in range(1, 20 + 1):
        print(f'--------------------正在打印第{page}页数据--------------------')
        url = f'https://www.csflhjw.com/zhenghun/9.html?page={page}'

        headers = {
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/92.0.4515.40 Safari/537.36'
        }

        resp = requests.get(url, headers = headers)
        if resp.status_code == 200:
            html_data = etree.HTML(resp.text)
            divs = html_data.xpath("//div[@class='zh-item']/div[@class='e']")


            for div in divs:
                # 获取当前页面信息
                name = div.xpath("./div[@class='e-name']/h2/text()")  # 小姐姐称呼
                name = ''.join(name)

                infos = div.xpath(".//div[@class='e-intro']/p[1]/text()")
                bir_date = ''.join(infos).split(' ')[0]  # 小姐姐出生日期
                height = ''.join(infos).split(' ')[1]  # 小姐姐身高
                educ_bgd = ''.join(infos).split(' ')[2]  # 小姐姐学历

                mary_stus = div.xpath(".//div[@class='e-intro']/p[2]/text()")  # 小姐姐婚否
                mary_stus = ''.join(mary_stus).split('：')[1]

                profe = div.xpath(".//div[@class='e-intro']/p[3]/text()")  # 小姐姐职业
                profe = ''.join(profe).split('：')[1]


                # 获取详情页面信息
                links = div.xpath("./div[@class='e-name']/a[@class='e-a']/@href")
                x = 'https://www.csflhjw.com'
                next_page = [x + i for i in links]


                for next in next_page:
                    response = requests.get(next, headers=headers)
                    if response.status_code == 200:
                        html = etree.HTML(response.text)

                        prof_phot = html.xpath("//div[@class='team-img']/img/@src")     # 小姐姐照片
                        x = 'https://www.csflhjw.com'
                        prof_phot = [x + i for i in prof_phot]
                        prof_phot = ''.join(prof_phot)

                        children = html.xpath("//div[@class='team-e']/p[4]/text()")     # 有无子女
                        children = ''.join(children).split('：')[1]

                        room = html.xpath("//div[@class='team-e']/p[5]/text()")         # 是否有房
                        room = ''.join(room).split('：')[1]

                        man_age = html.xpath("//div[@class='hunyin-1-2']/p[1]/span[1]/text()")     # 择偶年龄
                        man_age = ''.join(man_age).split('龄：')[1]

                        man_city = html.xpath("//div[@class='hunyin-1-2']/p[1]/span[2]/text()")        # 择偶城市
                        man_city = ''.join(man_city).split('市：')[1]

                        man_req = html.xpath("//div[@class='hunyin-1-2']/p[2]/span/text()")         # 男方要求
                        man_req = ''.join(man_req).split('求：')[1]

                        introd_myself = html.xpath("//div[@class='hunyin-1-3']/p/text()")       # 个人独白
                        introd_myself = ''.join(introd_myself).strip()

                        for paragraph in document.paragraphs:
                            # 段落对齐方式
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            # 行间距 注意用浮点数  2.0就表示两倍行间距
                            paragraph.paragraph_format.line_spacing = 2.0
                            # 段前与段后间距  Pt(12)表示12磅
                            paragraph.paragraph_format.space_before = Pt(12)  # 段前间距
                            paragraph.paragraph_format.space_after = Pt(12)  # 段后间距


                        document.add_heading(f'{count}号小姐姐来袭~~~').bold = True
                        document.add_paragraph(name)
                        document.add_paragraph(bir_date)
                        document.add_paragraph(height)
                        document.add_paragraph(educ_bgd)
                        document.add_paragraph(mary_stus)
                        document.add_paragraph(children)
                        document.add_paragraph(room)
                        document.add_paragraph(man_age)
                        document.add_paragraph(man_city)
                        document.add_paragraph(man_req)
                        document.add_paragraph(introd_myself)
                        document.add_page_break()  # 换页
                        count += 1

                        print(prof_phot,name, bir_date, height, educ_bgd, mary_stus, profe, children, room,
                              man_age, man_city, man_req, introd_myself)

    document.save('相亲网小姐姐.docx')


if __name__ == '__main__':
    spider_page()

